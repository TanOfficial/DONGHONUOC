import os
import re
import sys

# Reconfigure stdout to support UTF-8 encoding in Windows terminals
sys.stdout.reconfigure(encoding='utf-8')

def compile_markdown_to_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn
    except ImportError:
        print("[-] 'python-docx' library is not installed. Please run: pip install python-docx")
        sys.exit(1)

    print("[*] Initializing Word Document compiler...")
    doc = Document()


    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Configure Normal Style: Times New Roman 13pt
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Dark grey for academic readability
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # Helper function to add customized headings
    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy blue
            p.paragraph_format.space_before = Pt(24)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2E, 0x5B, 0x88) # Steel blue
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Dark grey
        else:
            # Level 4
            run.font.size = Pt(13)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55) # Medium grey
            p.paragraph_format.space_before = Pt(10)
            
        return p

    # Helper function to apply inline formatting (bold/italic) to a paragraph
    def add_formatted_text(paragraph, text, is_code=False):
        if is_code:
            run = paragraph.add_run(text)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xA3, 0x15, 0x15)
            return

        # Simple regex tokenizer for bold (**) and italic (*)
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for token in tokens:
            if token.startswith('**') and token.endswith('**'):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('*') and token.endswith('*'):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                paragraph.add_run(token)

    # Helper to color table cell background
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Helper to set cell borders thin
    def set_cell_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    # Chapter file paths
    base_dir = r"d:\ThucTap\DONGHONUOC\DONGHONUOC"
    chapters = [
        "chapter_1_job_background.md",
        "chapter_2_accomplishments.md",
        "chapter_3_technical_detail.md",
        "chapter_4_lessons_learned.md",
        "chapter_5_references.md"
    ]

    for idx, chap in enumerate(chapters):
        file_path = os.path.join(base_dir, chap)
        if not os.path.exists(file_path):
            print(f"[!] Chapter file not found: {chap}. Skipping...")
            continue

        print(f"[i] Compiling: {chap}...")
        with open(file_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        # State controllers
        in_code_block = False
        code_content = []
        table_rows = []
        in_table = False

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Handle Markdown Code Block boundaries
            if stripped.startswith("```"):
                if in_code_block:
                    # Closing code block: flush to word
                    in_code_block = False
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.line_spacing = 1.0
                    
                    # Add border decoration XML to paragraph
                    pPr = p._p.get_or_add_pPr()
                    pBdr = parse_xml(
                        f'<w:pBdr {nsdecls("w")}>\n'
                        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="888888"/>\n'
                        f'</w:pBdr>'
                    )
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F6"/>')
                    pPr.append(pBdr)
                    pPr.append(shd)

                    code_text = "".join(code_content)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    code_content = []
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_content.append(line)
                i += 1
                continue

            # Handle Markdown Tables
            if stripped.startswith("|"):
                in_table = True
                table_rows.append(stripped)
                i += 1
                continue
            elif in_table:
                # Table ended, compile it
                in_table = False
                
                # Filter out formatting rows like `| :--- | :--- |`
                clean_rows = []
                for tr in table_rows:
                    if re.match(r'^\|\s*[:\-|\s]+\s*\|$', tr):
                        continue
                    clean_rows.append(tr)
                
                if clean_rows:
                    # Extract cells
                    grid = []
                    for cr in clean_rows:
                        # Split by '|' and trim cells, ignoring outer splits
                        cells = [c.strip() for c in cr.split("|")[1:-1]]
                        grid.append(cells)
                    
                    if grid:
                        num_cols = len(grid[0])
                        w_table = doc.add_table(rows=len(grid), cols=num_cols)
                        w_table.autofit = True
                        
                        for row_idx, cells in enumerate(grid):
                            w_row = w_table.rows[row_idx]
                            is_header = (row_idx == 0)
                            
                            for col_idx, cell_text in enumerate(cells):
                                if col_idx < len(w_row.cells):
                                    w_cell = w_row.cells[col_idx]
                                    w_cell.text = "" # Clear default
                                    p = w_cell.paragraphs[0]
                                    p.paragraph_format.space_after = Pt(2)
                                    p.paragraph_format.line_spacing = 1.0
                                    
                                    run = p.add_run(cell_text)
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(11)
                                    
                                    set_cell_borders(w_cell)
                                    
                                    if is_header:
                                        run.bold = True
                                        set_cell_background(w_cell, "1B365D") # Navy blue header
                                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White text
                                    elif row_idx % 2 == 0:
                                        set_cell_background(w_cell, "F7F9FB") # Zebra striping
                
                table_rows = []
                # Fall through to process current line

            # Handle Alert Callouts (e.g. `> [!TIP]`)
            if stripped.startswith("> [!"):
                # Extract type
                alert_type = re.search(r'\[!(.*?)\]', stripped).group(1)
                i += 1
                desc_line = lines[i].strip()
                if desc_line.startswith(">"):
                    desc_line = desc_line.replace(">", "").strip()
                
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
                
                pPr = p._p.get_or_add_pPr()
                color_hex = "1B365D" # Default Navy
                bg_hex = "F0F4F8"
                if "WARNING" in alert_type or "CAUTION" in alert_type:
                    color_hex = "B71C1C" # Dark Red
                    bg_hex = "FFEBEE"
                elif "TIP" in alert_type:
                    color_hex = "2E7D32" # Dark Green
                    bg_hex = "E8F5E9"

                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>\n'
                    f'  <w:left w:val="single" w:sz="36" w:space="12" w:color="{color_hex}"/>\n'
                    f'</w:pBdr>'
                )
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_hex}"/>')
                pPr.append(pBdr)
                pPr.append(shd)
                
                # Add alert tag label
                run_tag = p.add_run(f"[{alert_type.upper()}] ")
                run_tag.bold = True
                run_tag.font.color.rgb = RGBColor(int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16))
                
                add_formatted_text(p, desc_line)
                i += 1
                continue

            # Standard Headings
            if stripped.startswith("#### "):
                add_custom_heading(stripped[5:], 4)
            elif stripped.startswith("### "):
                add_custom_heading(stripped[4:], 3)
            elif stripped.startswith("## "):
                add_custom_heading(stripped[3:], 2)
            elif stripped.startswith("# "):
                add_custom_heading(stripped[2:], 1)

            # Bullet points
            elif stripped.startswith("* ") or stripped.startswith("- "):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(4)
                add_formatted_text(p, stripped[2:])
            
            # Numbered Lists
            elif re.match(r'^\d+\.\s', stripped):
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_after = Pt(4)
                content = re.sub(r'^\d+\.\s', '', stripped)
                add_formatted_text(p, content)

            # Empty space lines
            elif not stripped:
                pass # Ignored to prevent excess page breaks/spacing gaps

            # Normal styled paragraphs
            else:
                p = doc.add_paragraph()
                add_formatted_text(p, stripped)

            i += 1

        # Add a Page Break between chapters (except for final chapter)
        if idx < len(chapters) - 1:
            doc.add_page_break()

    # Save final compiled document
    output_filename = os.path.join(base_dir, "BÁO_CÁO_THỰC_TẬP_HOÀN_CHỈNH.docx")
    doc.save(output_filename)
    print(f"[SUCCESS] Document successfully generated: {output_filename}")

if __name__ == "__main__":
    compile_markdown_to_docx()
